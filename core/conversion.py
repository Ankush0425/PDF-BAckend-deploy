"""
Conversion module — PDF to/from Word, Image, Text, Excel.
"""

import io
from pathlib import Path

import fitz  # PyMuPDF
from PIL import Image
from pypdf import PdfReader
import pdfplumber

from utils.file_manager import get_temp_path, get_original_filename
from utils.logger import get_logger

logger = get_logger("conversion")


# ──────────────────────────────────────────────
# PDF → Word
# ──────────────────────────────────────────────
def pdf_to_word(pdf_path: Path) -> Path:
    """Convert PDF to Word document (.docx) with high-fidelity formatting."""
    orig_name = get_original_filename(pdf_path)
    docx_name = Path(orig_name).with_suffix(".docx").name
    output_path = get_temp_path(docx_name)
    try:
        # Try pdf2docx for layout preservation
        try:
            logger.info("Using pdf2docx for layout-preserving PDF to DOCX conversion")
            from pdf2docx import Converter
            cv = Converter(str(pdf_path))
            cv.convert(str(output_path), start=0, end=None)
            cv.close()
            logger.info(f"Successfully converted PDF to Word via pdf2docx: {output_path.name}")
            return output_path
        except Exception as e:
            logger.warning(f"pdf2docx conversion failed, falling back to basic layout parser: {e}")

        # Fallback basic text/image parser if pdf2docx is unavailable
        from docx import Document
        from docx.shared import Inches, Pt

        doc = fitz.open(str(pdf_path))
        word_doc = Document()

        for i, page in enumerate(doc):
            if i > 0:
                word_doc.add_page_break()

            # Using dict format to preserve basic styles
            page_dict = page.get_text("dict")
            for block in page_dict.get("blocks", []):
                if block.get("type") == 0:  # Text block
                    for line in block.get("lines", []):
                        para = word_doc.add_paragraph()
                        for span in line.get("spans", []):
                            span_text = span.get("text", "")
                            if not span_text.strip():
                                continue
                            run = para.add_run(span_text)
                            
                            # Font size
                            size = span.get("size", 11)
                            run.font.size = Pt(size)
                            
                            # Bold & Italic checks
                            flags = span.get("flags", 0)
                            if flags & 2:  # Italic
                                run.italic = True
                            if flags & 16:  # Bold
                                run.bold = True
                            
                            font_name = span.get("font", "").lower()
                            if "bold" in font_name:
                                run.bold = True
                            if "italic" in font_name or "oblique" in font_name:
                                run.italic = True

            # Extract and add images
            images = page.get_images(full=True)
            for img_idx, img_info in enumerate(images):
                try:
                    xref = img_info[0]
                    base_image = doc.extract_image(xref)
                    img_bytes = base_image["image"]
                    img_ext = base_image["ext"]

                    img_path = get_temp_path(f"img_{i}_{img_idx}.{img_ext}")
                    img_path.write_bytes(img_bytes)
                    word_doc.add_picture(str(img_path), width=Inches(5))
                except Exception:
                    pass

        word_doc.save(str(output_path))
        doc.close()
        logger.info(f"Converted PDF to Word (basic fallback): {output_path.name}")
        return output_path
    except Exception as e:
        logger.error(f"PDF to Word failed: {e}")
        raise RuntimeError(f"PDF to Word conversion failed: {e}") from e


# ──────────────────────────────────────────────
# Word → PDF
# ──────────────────────────────────────────────
def word_to_pdf(docx_path: Path) -> Path:
    """Convert Word document to PDF using LibreOffice (best fidelity) or ReportLab (fallback)."""
    orig_name = get_original_filename(docx_path)
    pdf_name = Path(orig_name).with_suffix(".pdf").name
    output_path = get_temp_path(pdf_name)
    
    # Try LibreOffice headless for high-fidelity conversion
    import subprocess
    import shutil
    
    soffice_path = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice_path:
        try:
            logger.info("Using LibreOffice headless for high-fidelity DOCX to PDF conversion")
            cmd = [
                soffice_path,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(output_path.parent),
                str(docx_path)
            ]
            result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
            
            # LibreOffice generates output with the same base name as docx_path.
            generated_pdf = docx_path.parent / docx_path.with_suffix(".pdf").name
            if generated_pdf.exists():
                shutil.move(str(generated_pdf), str(output_path))
                logger.info(f"Successfully converted Word to PDF via LibreOffice: {output_path.name}")
                return output_path
        except Exception as e:
            logger.warning(f"LibreOffice conversion failed, falling back to ReportLab: {e}")

    # Fallback to reportlab
    from docx import Document
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT

    try:
        word_doc = Document(str(docx_path))
        pdf_doc = SimpleDocTemplate(str(output_path), pagesize=A4,
                                     leftMargin=inch, rightMargin=inch,
                                     topMargin=inch, bottomMargin=inch)

        styles = getSampleStyleSheet()
        body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=11,
                                     leading=14, alignment=TA_LEFT)
        heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading1'],
                                        fontSize=16, leading=20, spaceAfter=12)

        story = []
        for para in word_doc.paragraphs:
            para_html = ""
            for run in para.runs:
                run_text = run.text
                if not run_text:
                    continue
                # Escape special characters for ReportLab
                run_text = run_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                # Apply styles
                if run.bold:
                    run_text = f"<b>{run_text}</b>"
                if run.italic:
                    run_text = f"<i>{run_text}</i>"
                if run.font.size:
                    pt_size = run.font.size.pt
                    run_text = f'<font size="{pt_size}">{run_text}</font>'
                if run.font.color and run.font.color.rgb:
                    color_hex = f"#{run.font.color.rgb}"
                    run_text = f'<font color="{color_hex}">{run_text}</font>'
                
                para_html += run_text

            if not para_html.strip():
                story.append(Spacer(1, 6))
                continue

            if para.style.name.startswith('Heading'):
                story.append(Paragraph(para_html, heading_style))
            else:
                story.append(Paragraph(para_html, body_style))

        if not story:
            story.append(Paragraph("(Empty document)", body_style))

        pdf_doc.build(story)
        logger.info(f"Converted Word to PDF (ReportLab fallback): {output_path.name}")
        return output_path
    except Exception as e:
        logger.error(f"Word to PDF failed: {e}")
        raise RuntimeError(f"Word to PDF conversion failed: {e}") from e


# ──────────────────────────────────────────────
# PDF → Image
# ──────────────────────────────────────────────
def pdf_to_images(pdf_path: Path, dpi: int = 150, fmt: str = "png") -> list[Path]:
    """Convert each PDF page to an image."""
    output_paths = []
    try:
        doc = fitz.open(str(pdf_path))
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)

        for i, page in enumerate(doc):
            pix = page.get_pixmap(matrix=mat)
            out_path = get_temp_path(f"page_{i + 1}.{fmt}")
            if fmt.lower() == "jpg" or fmt.lower() == "jpeg":
                pix.save(str(out_path), jpg_quality=90)
            else:
                pix.save(str(out_path))
            output_paths.append(out_path)

        doc.close()
        logger.info(f"Converted PDF to {len(output_paths)} images")
        return output_paths
    except Exception as e:
        logger.error(f"PDF to images failed: {e}")
        raise RuntimeError(f"PDF to Image conversion failed: {e}") from e


# ──────────────────────────────────────────────
# Image → PDF
# ──────────────────────────────────────────────
def image_to_pdf(image_path: Path) -> Path:
    """Convert a single image to PDF."""
    orig_name = get_original_filename(image_path)
    pdf_name = Path(orig_name).with_suffix(".pdf").name
    output_path = get_temp_path(pdf_name)
    try:
        img = Image.open(str(image_path))
        if img.mode == 'RGBA':
            img = img.convert('RGB')
        img.save(str(output_path), "PDF", resolution=150)
        logger.info(f"Converted image to PDF: {output_path.name}")
        return output_path
    except Exception as e:
        logger.error(f"Image to PDF failed: {e}")
        raise RuntimeError(f"Image to PDF conversion failed: {e}") from e


# ──────────────────────────────────────────────
# PDF → Text
# ──────────────────────────────────────────────
def pdf_to_text(pdf_path: Path) -> tuple[Path, str]:
    """Extract text from PDF and save to .txt file."""
    orig_name = get_original_filename(pdf_path)
    txt_name = Path(orig_name).with_suffix(".txt").name
    output_path = get_temp_path(txt_name)
    try:
        doc = fitz.open(str(pdf_path))
        text_parts = []
        for i, page in enumerate(doc):
            text = page.get_text()
            text_parts.append(f"--- Page {i + 1} ---\n{text}")
        doc.close()

        full_text = "\n\n".join(text_parts)
        output_path.write_text(full_text, encoding="utf-8")
        logger.info(f"Extracted text from PDF ({len(full_text)} chars)")
        return output_path, full_text
    except Exception as e:
        logger.error(f"PDF to text failed: {e}")
        raise RuntimeError(f"PDF to Text conversion failed: {e}") from e


# ──────────────────────────────────────────────
# Text → PDF
# ──────────────────────────────────────────────
def text_to_pdf(text: str, output_name: str = "from_text.pdf") -> Path:
    """Convert plain text to a PDF document."""
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch

    output_path = get_temp_path(output_name)
    try:
        doc = SimpleDocTemplate(str(output_path), pagesize=A4,
                                 leftMargin=inch, rightMargin=inch,
                                 topMargin=inch, bottomMargin=inch)
        styles = getSampleStyleSheet()
        story = []

        for line in text.split('\n'):
            escaped = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            if escaped.strip():
                story.append(Paragraph(escaped, styles['Normal']))
            else:
                story.append(Spacer(1, 6))

        if not story:
            story.append(Paragraph("(Empty document)", styles['Normal']))

        doc.build(story)
        logger.info(f"Converted text to PDF: {output_path.name}")
        return output_path
    except Exception as e:
        logger.error(f"Text to PDF failed: {e}")
        raise RuntimeError(f"Text to PDF conversion failed: {e}") from e


# ──────────────────────────────────────────────
# Excel → PDF
# ──────────────────────────────────────────────
def excel_to_pdf(excel_path: Path) -> Path:
    """Convert Excel spreadsheet to PDF."""
    from openpyxl import load_workbook
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    from reportlab.lib.units import inch

    orig_name = get_original_filename(excel_path)
    pdf_name = Path(orig_name).with_suffix(".pdf").name
    output_path = get_temp_path(pdf_name)
    try:
        wb = load_workbook(str(excel_path), data_only=True)
        doc = SimpleDocTemplate(str(output_path), pagesize=landscape(A4),
                                 leftMargin=0.5*inch, rightMargin=0.5*inch,
                                 topMargin=0.5*inch, bottomMargin=0.5*inch)
        styles = getSampleStyleSheet()
        story = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            story.append(Paragraph(f"Sheet: {sheet_name}", styles['Heading2']))
            story.append(Spacer(1, 12))

            data = []
            for row in ws.iter_rows(values_only=True):
                data.append([str(cell) if cell is not None else "" for cell in row])

            if data:
                table = Table(data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#6C63FF')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F5F6FA')]),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ]))
                story.append(table)
            story.append(Spacer(1, 24))

        doc.build(story)
        logger.info(f"Converted Excel to PDF: {output_path.name}")
        return output_path
    except Exception as e:
        logger.error(f"Excel to PDF failed: {e}")
        raise RuntimeError(f"Excel to PDF conversion failed: {e}") from e


# ──────────────────────────────────────────────
# PDF → Excel
# ──────────────────────────────────────────────
def pdf_to_excel(pdf_path: Path) -> Path:
    """Extract tables from PDF to Excel."""
    from openpyxl import Workbook

    orig_name = get_original_filename(pdf_path)
    xlsx_name = Path(orig_name).with_suffix(".xlsx").name
    output_path = get_temp_path(xlsx_name)
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Extracted Data"

        with pdfplumber.open(str(pdf_path)) as pdf:
            current_row = 1
            for i, page in enumerate(pdf.pages):
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        for row in table:
                            for col_idx, cell in enumerate(row):
                                ws.cell(row=current_row, column=col_idx + 1, value=cell or "")
                            current_row += 1
                        current_row += 1  # Blank row between tables
                else:
                    # If no tables, extract text
                    text = page.extract_text()
                    if text:
                        ws.cell(row=current_row, column=1, value=f"Page {i + 1}: {text[:1000]}")
                        current_row += 1

        wb.save(str(output_path))
        logger.info(f"Extracted tables to Excel: {output_path.name}")
        return output_path
    except Exception as e:
        logger.error(f"PDF to Excel failed: {e}")
        raise RuntimeError(f"PDF to Excel conversion failed: {e}") from e
